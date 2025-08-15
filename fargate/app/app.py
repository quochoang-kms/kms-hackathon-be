from fastapi import FastAPI, HTTPException, File, UploadFile, Form
from fastapi.responses import PlainTextResponse, StreamingResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List, Union
import os
import sys
import asyncio
import json
import logging
from dotenv import load_dotenv
import uvicorn
import PyPDF2
import docx
import io
import re
import time
import traceback

from strands import Agent

from agents.jd_analyzer.jd_analyzer import jd_analyzer
from agents.cv_analyzer.cv_analyzer import cv_analyzer  
from agents.skill_matcher.skill_matcher import skill_matcher
from agents.question_generator.question_generator import question_generator
from models import JDResponse, CVResponse, SkillMatcherResponse, QuestionGeneratorResponse
from conditions.conditions import is_matched_skill, is_analyzer_done, is_skill_matching_done

from strands.multiagent import GraphBuilder
from strands.models import BedrockModel
from strands.types.content import ContentBlock

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Add the agents path to sys.path
sys.path.append('/app')

# Initialize FastAPI app
app = FastAPI(
    title="Interview Preparation API",
    description="AI-powered interview preparation system using Strands agents",
    version="1.0.0"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Bedrock Model Config
bedrock_model = BedrockModel(
    model_id=os.getenv("MODEL_ID"),
    region_name=os.getenv("REGION_NAME"),
)

# Initialize the multi-agent graph
def create_interview_graph():
    """Create and configure the interview preparation agent graph"""
        
    ORCHESTRATOR_PROMPT = """
    You are the ORCHESTRATOR, a routing agent that coordinates multiple specialized AI agents in an interview preparation system.

    Your role is to:
    1. Analyze incoming requests and documents (JD, CV)
    2. Route tasks to appropriate specialized agents:
       - JD_ANALYZER: For job description analysis
       - CV_ANALYZER: For candidate CV/resume analysis
       - SKILL_MATCHER: For matching skills between JD and CV
       - QUESTION_GENERATOR: For generating tailored interview questions
    3. Coordinate the workflow between agents
    4. Ensure proper data flow and task completion

    When you receive documents, determine which agents need to process them and route accordingly.
    Always maintain context and ensure all necessary agents are activated for comprehensive analysis.
    The workflow should be: JD_ANALYZER & CV_ANALYZER → SKILL_MATCHER → QUESTION_GENERATOR
    """
    
    orchestrator = Agent(
        name="ORCHESTRATOR",
        model=bedrock_model,
        system_prompt=ORCHESTRATOR_PROMPT,
    )

    # Initialize GraphBuilder
    builder = GraphBuilder()

    # Add Agent Nodes
    builder.add_node(orchestrator, "ORCHESTRATOR")
    builder.add_node(jd_analyzer, "JD_ANALYZER")
    builder.add_node(cv_analyzer, "CV_ANALYZER")
    builder.add_node(skill_matcher, "SKILL_MATCHER")
    builder.add_node(question_generator, "QUESTION_GENERATOR")

    # Add Edges
    builder.add_edge("ORCHESTRATOR", "JD_ANALYZER")
    builder.add_edge("ORCHESTRATOR", "CV_ANALYZER")
    builder.add_edge("JD_ANALYZER", "SKILL_MATCHER", condition=is_analyzer_done)
    builder.add_edge("CV_ANALYZER", "SKILL_MATCHER", condition=is_analyzer_done)
    builder.add_edge("SKILL_MATCHER", "QUESTION_GENERATOR", condition=is_skill_matching_done)

    # Define the entry point for the orchestrator
    builder.set_entry_point("ORCHESTRATOR")

    return builder.build()

# Global graph instance
interview_graph = create_interview_graph()

# Create streaming-enabled orchestrator agent
def create_streaming_orchestrator():
    """Create orchestrator agent with streaming capabilities"""
    ORCHESTRATOR_PROMPT = """
    You are an AI assistant specialized in interview preparation. Analyze the provided Job Description and CV, then provide comprehensive interview preparation including:
    
    1. Job Description Analysis
    2. CV Analysis 
    3. Skills Matching Assessment
    4. Tailored Interview Questions
    
    Format your response as structured JSON with clear sections for each analysis.
    """
    
    return Agent(
        name="STREAMING_ORCHESTRATOR",
        model=bedrock_model,
        system_prompt=ORCHESTRATOR_PROMPT,
        callback_handler=None  # Disable default callback for streaming
    )

# Streaming generator function using Strands native streaming
async def stream_interview_process(processed_jd_text: str, processed_cv_text: str):
    """Generator function that yields real-time streaming updates from Strands agents"""
    start_time = time.time()
    
    try:
        # Send start event
        yield f"data: {json.dumps({'event': 'started', 'data': {'timestamp': time.strftime('%Y-%m-%d %H:%M:%S'), 'progress': 0}})}\n\n"
        
        # Create streaming orchestrator
        streaming_orchestrator = create_streaming_orchestrator()
        
        # Prepare the query with both JD and CV content
        query = f"""
        Please analyze the following Job Description and CV, then provide comprehensive interview preparation:
        
        JOB DESCRIPTION:
        {processed_jd_text}
        
        CV/RESUME:
        {processed_cv_text}
        
        Please provide:
        1. Job Description Analysis
        2. CV Analysis 
        3. Skills Matching Assessment
        4. Tailored Interview Questions
        
        Format your response as structured JSON with clear sections for each analysis.
        """
        
        # Send processing event
        yield f"data: {json.dumps({'event': 'processing', 'data': {'message': 'Starting AI analysis', 'progress': 10}})}\n\n"
        
        # Stream the agent's response using Strands native streaming
        agent_stream = streaming_orchestrator.stream_async(query)
        
        current_progress = 20
        response_chunks = []
        
        async for event in agent_stream:
            if "data" in event:
                # Stream text chunks as they're generated
                chunk = event["data"]
                response_chunks.append(chunk)
                
                # Send streaming text event
                yield f"data: {json.dumps({'event': 'text_chunk', 'data': {'chunk': chunk, 'progress': current_progress}})}\n\n"
                
                # Gradually increase progress
                current_progress = min(85, current_progress + 1)
                
            elif "current_tool_use" in event and event["current_tool_use"].get("name"):
                # Send tool usage event
                tool_name = event["current_tool_use"]["name"]
                yield f"data: {json.dumps({'event': 'tool_use', 'data': {'tool': tool_name, 'progress': current_progress}})}\n\n"
        
        # Combine all response chunks
        full_response = ''.join(response_chunks)
        
        # Send processing completion
        yield f"data: {json.dumps({'event': 'processing_complete', 'data': {'message': 'Analysis complete, parsing results', 'progress': 90}})}\n\n"
        
        # Parse the response for structured data
        try:
            # Try to extract JSON from the response
            json_match = re.search(r'```json\s*(.*?)\s*```', full_response, re.DOTALL)
            if json_match:
                parsed_result = json.loads(json_match.group(1))
            else:
                # If no JSON block found, create a structured response
                parsed_result = {
                    "analysis": full_response,
                    "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
                }
        except json.JSONDecodeError:
            parsed_result = {
                "analysis": full_response,
                "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
            }
        
        # Send final result
        final_result = {
            "status": "completed",
            "execution_time": time.time() - start_time,
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
            "result": parsed_result,
            "raw_response": full_response
        }
        
        yield f"data: {json.dumps({'event': 'completed', 'data': final_result, 'progress': 100})}\n\n"
        
    except Exception as e:
        error_data = {
            "error": str(e),
            "execution_time": time.time() - start_time,
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
        }
        yield f"data: {json.dumps({'event': 'error', 'data': error_data})}\n\n"

# Pydantic models for request/response
class InterviewRequest(BaseModel):
    job_description: Optional[str] = None
    cv_content: Optional[str] = None
    additional_context: Optional[str] = None

class InterviewResponse(BaseModel):
    status: str
    execution_time: Optional[float] = None
    jd_analysis: Optional[JDResponse] = None
    cv_analysis: Optional[CVResponse] = None
    skills_matching: Optional[SkillMatcherResponse] = None
    questions: Optional[QuestionGeneratorResponse] = None
    summary: Optional[dict] = None

class HealthResponse(BaseModel):
    status: str
    service: str
    version: str

class StreamEvent(BaseModel):
    event: str
    data: dict
    timestamp: Optional[str] = None

# Health check endpoint
@app.get("/health", response_model=HealthResponse)
async def health_check():
    """Health check endpoint for load balancer"""
    return HealthResponse(
        status="healthy",
        service="interview-preparation-api",
        version="1.0.0"
    )


# File upload endpoint for handling CV/JD uploads
@app.post("/prepare-interview")
async def prepare_interview(
    jd_text: str = Form(...),
    cv_file: UploadFile = File(...),
):
    """
    Prepare interview questions using JD text and CV file
    JD: Text input, CV: PDF or DOCX file
    """
    start_time = time.time()
    
    try:
        # Function to extract text from PDF
        def extract_pdf_text(content: bytes, filename: str) -> str:
            try:
                pdf_file = io.BytesIO(content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
            except Exception as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unable to extract text from PDF {filename}: {str(e)}"
                )
        
        # Function to extract text from DOCX
        def extract_docx_text(content: bytes, filename: str) -> str:
            try:
                doc_file = io.BytesIO(content)
                doc = docx.Document(doc_file)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text.strip()
            except Exception as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unable to extract text from DOCX {filename}: {str(e)}"
                )
        

        # Function to process file based on type
        def process_file(content: bytes, file: UploadFile) -> str:
            filename = file.filename.lower()
            content_type = file.content_type
            
            if content_type == 'application/pdf' or filename.endswith('.pdf'):
                return extract_pdf_text(content, file.filename)
            elif (content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or 
                  filename.endswith('.docx')):
                return extract_docx_text(content, file.filename)
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {content_type}. Please upload PDF or DOCX files only."
                )
        
        # Process inputs
        processed_jd_text = jd_text.strip()
        
        cv_content = await cv_file.read()
        processed_cv_text = process_file(cv_content, cv_file)
        
        # Additional validation - check if content makes sense
        if len(processed_jd_text.strip()) < 10:
            raise HTTPException(status_code=400, detail="Job description appears to be too short or invalid")
        
        if len(processed_cv_text.strip()) < 10:
            raise HTTPException(status_code=400, detail="CV content appears to be too short or invalid")
        
        # Create content blocks for the graph
        content_blocks = [
            ContentBlock(text="Start Interview Preparation System"),
            ContentBlock(
                document={
                    "name": "JD",
                    "format": "txt",
                    "source": {
                        "bytes": processed_jd_text.encode('utf-8'),
                    },
                }
            ),
            ContentBlock(
                document={
                    "name": "CV",
                    "format": "txt",
                    "source": {
                        "bytes": processed_cv_text.encode('utf-8'),
                    },
                }
            ),
        ]
        
        # Execute the agent graph
        logger.info("Executing multi-agent workflow")
        result = interview_graph(content_blocks)
        
        jd_analyzer_response = result.results["JD_ANALYZER"].result.message["content"][0]["text"]
        cv_analyzer_response = result.results["CV_ANALYZER"].result.message["content"][0]["text"]
        
        skill_matcher_response = result.results["SKILL_MATCHER"].result.message["content"][0]["text"]
        question_generator_response = result.results["QUESTION_GENERATOR"].result.message["content"][0]["text"]
        
        jd_analyzer_parser = re.search(r'```json\s*(.*?)\s*```', jd_analyzer_response, re.DOTALL)
        cv_analyzer_parser = re.search(r'```json\s*(.*?)\s*```', cv_analyzer_response, re.DOTALL)
        skill_match_parser = re.search(r'```json\s*(.*?)\s*```', skill_matcher_response, re.DOTALL)
        question_generator_parser = re.search(r'```json\s*(.*?)\s*```', question_generator_response, re.DOTALL)
        
        if jd_analyzer_parser:
            jd_analyzer_json = json.loads(jd_analyzer_parser.group(1))
        
        if cv_analyzer_parser:
            cv_analyzer_json = json.loads(cv_analyzer_parser.group(1))
        
        if skill_match_parser:
            skill_matcher_json = json.loads(skill_match_parser.group(1))
            
        if question_generator_parser:
            question_generator_json = json.loads(question_generator_parser.group(1))
       
        response_data = {
            "status": "completed",
            "execution_time": time.time() - start_time,
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
            "jd_analysis": jd_analyzer_json["JDResponse"],
            "cv_analysis": cv_analyzer_json["CVResponse"],
            "skill_matcher": skill_matcher_json["SkillMatcherResponse"],
            "question_generator": question_generator_json["QuestionGeneratorResponse"],            
        }
        
        return JSONResponse(
            content=response_data,
            headers={
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "POST, OPTIONS",
                "Access-Control-Allow-Headers": "*",
            }
        )
    
    except HTTPException:
        # Re-raise HTTP exceptions (validation errors)
        raise
    except Exception as e:
        execution_time = time.time() - start_time
        logger.error(f"File processing failed after {execution_time:.2f} seconds: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        
        # Return error response instead of raising exception
        error_data = {
            "status": "error",
            "execution_time": execution_time,
            "error": str(e),
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
        }
        
        return JSONResponse(
            content=error_data,
            status_code=500,
            headers={
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Methods": "POST, OPTIONS",
                "Access-Control-Allow-Headers": "*",
            }
        )


# Streaming endpoint
@app.post("/prepare-interview-stream")
async def prepare_interview_stream(
    jd_text: str = Form(...),
    cv_file: UploadFile = File(...),
):
    """
    Prepare interview questions with streaming updates
    JD: Text input, CV: PDF or DOCX file
    Returns Server-Sent Events stream
    
    Required form fields:
    - jd_text: Job description as text
    - cv_file: CV file (PDF or DOCX)
    """
    try:
        # Validate inputs exist
        if not jd_text or not jd_text.strip():
            raise HTTPException(status_code=400, detail="jd_text is required")
        
        if not cv_file or not cv_file.filename:
            raise HTTPException(status_code=400, detail="cv_file is required")
        # Function to extract text from PDF
        def extract_pdf_text(content: bytes, filename: str) -> str:
            try:
                pdf_file = io.BytesIO(content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
            except Exception as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unable to extract text from PDF {filename}: {str(e)}"
                )
        
        # Function to extract text from DOCX
        def extract_docx_text(content: bytes, filename: str) -> str:
            try:
                doc_file = io.BytesIO(content)
                doc = docx.Document(doc_file)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text.strip()
            except Exception as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unable to extract text from DOCX {filename}: {str(e)}"
                )
        
        # Function to process file based on type
        def process_file(content: bytes, file: UploadFile) -> str:
            filename = file.filename.lower()
            content_type = file.content_type
            
            if content_type == 'application/pdf' or filename.endswith('.pdf'):
                return extract_pdf_text(content, file.filename)
            elif (content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or 
                  filename.endswith('.docx')):
                return extract_docx_text(content, file.filename)
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {content_type}. Please upload PDF or DOCX files only."
                )
        
        # Process inputs
        processed_jd_text = jd_text.strip()
        cv_content = await cv_file.read()
        processed_cv_text = process_file(cv_content, cv_file)
        
        # Validation
        if len(processed_jd_text.strip()) < 10:
            raise HTTPException(status_code=400, detail="Job description appears to be too short or invalid")
        
        if len(processed_cv_text.strip()) < 10:
            raise HTTPException(status_code=400, detail="CV content appears to be too short or invalid")
        
        # Return streaming response
        return StreamingResponse(
            stream_interview_process(processed_jd_text, processed_cv_text),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "Access-Control-Allow-Origin": "*",
                "Access-Control-Allow-Headers": "*",
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Streaming setup failed: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
