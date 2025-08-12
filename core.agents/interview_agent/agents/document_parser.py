"""Document Parser Agent for extracting text from PDF and DOCX files"""

import os
import asyncio
from typing import Union, Dict, Any
import PyPDF2
import pdfplumber
from docx import Document
from strands_agents import Agent, tool
import logging

logger = logging.getLogger(__name__)

class DocumentParserAgent(Agent):
    """Agent for parsing documents and extracting text content"""
    
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
    
    @tool
    async def parse_document(self, input_data: Union[str, bytes], input_type: str = "auto") -> Dict[str, Any]:
        """Parse document and extract text content
        
        Args:
            input_data: File path, text content, or binary data
            input_type: Type of input ("text", "file_path", "binary", "auto")
            
        Returns:
            Dict with extracted text and metadata
        """
        try:
            if input_type == "auto":
                input_type = self._detect_input_type(input_data)
            
            if input_type == "text":
                return await self._process_text_input(input_data)
            elif input_type == "file_path":
                return await self._process_file_input(input_data)
            elif input_type == "binary":
                return await self._process_binary_input(input_data)
            else:
                raise ValueError(f"Unsupported input type: {input_type}")
                
        except Exception as e:
            logger.error(f"Document parsing failed: {str(e)}")
            return {"text": "", "error": str(e), "metadata": {}}
    
    def _detect_input_type(self, input_data: Union[str, bytes]) -> str:
        """Detect the type of input data"""
        if isinstance(input_data, bytes):
            return "binary"
        elif isinstance(input_data, str):
            if os.path.isfile(input_data):
                return "file_path"
            else:
                return "text"
        return "text"
    
    async def _process_text_input(self, text: str) -> Dict[str, Any]:
        """Process direct text input"""
        cleaned_text = self._clean_text(text)
        return {
            "text": cleaned_text,
            "metadata": {
                "source": "direct_text",
                "length": len(cleaned_text),
                "word_count": len(cleaned_text.split())
            }
        }
    
    async def _process_file_input(self, file_path: str) -> Dict[str, Any]:
        """Process file input based on extension"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return await self._parse_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return await self._parse_docx(file_path)
        else:
            # Try to read as text file
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return await self._process_text_input(text)
    
    async def _process_binary_input(self, binary_data: bytes) -> Dict[str, Any]:
        """Process binary data input"""
        # For now, assume it's PDF binary data
        # In production, you'd detect the file type from binary headers
        try:
            import io
            pdf_file = io.BytesIO(binary_data)
            reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            cleaned_text = self._clean_text(text)
            return {
                "text": cleaned_text,
                "metadata": {
                    "source": "binary_pdf",
                    "pages": len(reader.pages),
                    "length": len(cleaned_text)
                }
            }
        except Exception as e:
            logger.error(f"Binary parsing failed: {str(e)}")
            return {"text": "", "error": str(e), "metadata": {}}
    
    async def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Parse PDF file using pdfplumber for better text extraction"""
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            cleaned_text = self._clean_text(text)
            return {
                "text": cleaned_text,
                "metadata": {
                    "source": file_path,
                    "pages": len(pdf.pages),
                    "length": len(cleaned_text)
                }
            }
        except Exception as e:
            logger.error(f"PDF parsing failed: {str(e)}")
            return {"text": "", "error": str(e), "metadata": {}}
    
    async def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parse DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            cleaned_text = self._clean_text(text)
            return {
                "text": cleaned_text,
                "metadata": {
                    "source": file_path,
                    "paragraphs": len(doc.paragraphs),
                    "length": len(cleaned_text)
                }
            }
        except Exception as e:
            logger.error(f"DOCX parsing failed: {str(e)}")
            return {"text": "", "error": str(e), "metadata": {}}
    
    def _clean_text(self, text: str) -> str:
        """Clean and preprocess extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        lines = [line for line in lines if line]
        
        # Join lines and normalize spacing
        cleaned = ' '.join(lines)
        
        # Remove multiple spaces
        import re
        cleaned = re.sub(r'\s+', ' ', cleaned)
        
        return cleaned.strip()