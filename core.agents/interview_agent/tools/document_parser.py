import os
import re
import sys
from typing import Optional

# Add parent directory to path for imports
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

try:
    from strands import tool
except ImportError:
    # Fallback decorator if strands not available
    def tool(func):
        return func

# Try to import DocumentContent, create fallback if not available
try:
    from base_models import DocumentContent
except ImportError:
    # Fallback DocumentContent class
    class DocumentContent:
        def __init__(self, content: str, file_type: str = "text"):
            self.content = content
            self.file_type = file_type

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None


@tool
def parse_pdf(file_path: str) -> DocumentContent:
    """
    Parse PDF file and extract text content.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        DocumentContent: Extracted content and metadata
    """
    if not PyPDF2:
        raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install PyPDF2")
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    content = ""
    metadata = {"pages": 0, "file_size": os.path.getsize(file_path)}
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            metadata["pages"] = len(pdf_reader.pages)
            
            for page in pdf_reader.pages:
                content += page.extract_text() + "\n"
                
        # Clean up the extracted text
        content = _clean_text(content)
        
        return DocumentContent(
            content=content,
            metadata=metadata,
            file_type="pdf"
        )
        
    except Exception as e:
        raise Exception(f"Error parsing PDF: {str(e)}")


@tool
def parse_docx(file_path: str) -> DocumentContent:
    """
    Parse DOCX file and extract text content.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        DocumentContent: Extracted content and metadata
    """
    if not Document:
        raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    try:
        doc = Document(file_path)
        content = ""
        
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += cell.text + " "
                content += "\n"
        
        # Clean up the extracted text
        content = _clean_text(content)
        
        metadata = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "file_size": os.path.getsize(file_path)
        }
        
        return DocumentContent(
            content=content,
            metadata=metadata,
            file_type="docx"
        )
        
    except Exception as e:
        raise Exception(f"Error parsing DOCX: {str(e)}")


@tool
def parse_text_file(file_path: str) -> DocumentContent:
    """
    Parse plain text file.
    
    Args:
        file_path: Path to the text file
        
    Returns:
        DocumentContent: Extracted content and metadata
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Clean up the text
        content = _clean_text(content)
        
        metadata = {
            "file_size": os.path.getsize(file_path),
            "lines": len(content.split('\n'))
        }
        
        return DocumentContent(
            content=content,
            metadata=metadata,
            file_type="txt"
        )
        
    except Exception as e:
        raise Exception(f"Error parsing text file: {str(e)}")


def _clean_text(text: str) -> str:
    """
    Clean and normalize extracted text.
    
    Args:
        text: Raw extracted text
        
    Returns:
        str: Cleaned text
    """
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    
    # Remove special characters that might interfere with processing
    text = re.sub(r'[^\w\s\-.,;:!?()[\]{}"\'/]', ' ', text)
    
    # Normalize line breaks
    text = re.sub(r'\n+', '\n', text)
    
    # Strip leading/trailing whitespace
    text = text.strip()
    
    return text
